# PEGA OS DADOS DE UMA PLANILHA, COLOCA EM UM DOCX E CONVERTE PARA UM PDF
import openpyxl
from docx import Document
from docx2pdf import convert
import os

# Função para ler palavras-chave da planilha
def read_keywords_from_excel(file_path, sheet_name='Semana4', column=1):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb[sheet_name]
    keywords = []

    for row in sheet.iter_rows(min_col=3, max_col=3, values_only=True):
        if row[0]:
            keywords.append(row[0])

    return keywords

# Função para adicionar palavras-chave ao documento Word
def create_doc_with_keywords(keywords, output_path):
    doc = Document()
    doc.add_heading('Qual o seu nome completo', level=1)

    for keyword in keywords:
        doc.add_paragraph(keyword)

    doc.save(output_path)
    print(f"Arquivo .docx criado: {output_path}")

# Função para converter o documento .docx para PDF
def convert_doc_to_pdf(input_path, output_path=None):
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + ".pdf"

    convert(input_path, output_path)
    print(f"Arquivo PDF criado: {output_path}")

excel_file = 'ODS - Entregas.xlsx'  # Planilha com palavras-chaves
docx_file = 'palavras_chaves.docx'   # Nome do arquivo .docx de saída

# 1. Ler palavras-chave da planilha
keywords = read_keywords_from_excel(excel_file)

# 2. Criar o arquivo .docx com as palavras-chave
create_doc_with_keywords(keywords, docx_file)

# 3. Converter o arquivo .docx para PDF
convert_doc_to_pdf(docx_file)
